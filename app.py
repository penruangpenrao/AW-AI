from flask import Flask, request, send_file, render_template, jsonify
import docx
from docx.shared import RGBColor
import io
import json
import zipfile
import os
import google.generativeai as genai

app = Flask(__name__)

# ── 1. ตั้งค่า Google Gemini API ──
# นำ API Key ฟรีของคุณจาก Google AI Studio มาใส่ที่นี่ครับ
GEMINI_API_KEY = "AIzaSyBhcoxnVIXRk-mLZUlqAuE3RZMQ5jxqEE0"
genai.configure(api_key=GEMINI_API_KEY)
# ใช้รุ่น flash เพราะทำงานไวและฟรีโควตาเยอะเหมาะกับงานเอกสาร
model = genai.GenerativeModel('gemini-2.5-flash')# ── CORS ──
@app.after_request
def add_cors(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

# ── ฟังก์ชันตัวช่วยจัดการ Word ──
def is_reddish(run):
    if run.font.color and run.font.color.rgb:
        hex_color = str(run.font.color.rgb).upper()
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            if r > 130 and g < 100 and b < 100:
                return True
        except:
            pass
    return False

def normalize_red_runs(doc):
    def _normalize_paragraph(p):
        first_red_run = None
        for run in p.runs:
            if is_reddish(run):
                if first_red_run is None:
                    first_red_run = run
                else:
                    first_red_run.text += run.text
                    run.text = ""
            else:
                first_red_run = None

    for p in doc.paragraphs:
        _normalize_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _normalize_paragraph(p)

def extract_all_red_words(doc):
    red_words = set()
    def _extract(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                if is_reddish(run):
                    text = run.text.strip()
                    if text:
                        red_words.add(text)
                        
    _extract(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _extract(cell.paragraphs)
    return list(red_words)

# ── Routes ──
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process_ai', methods=['POST'])
def process_ai():
    files = request.files.getlist('files')
    excel_data_str = request.form.get('excel_data')

    if not files or not excel_data_str:
        return jsonify({'error': 'ข้อมูลไม่ครบถ้วน'}), 400

    try:
        excel_data = json.loads(excel_data_str)
        
        # 1. อ่านไฟล์ทั้งหมดและรวบรวม "คำสีแดง" ทุกคำจากทุกไฟล์ เพื่อส่งให้ AI คิดทีเดียว
        all_unique_red_words = set()
        loaded_docs = [] # เก็บไฟล์ Word ที่โหลดไว้แล้วจะได้ไม่ต้องโหลดใหม่
        
        for file in files:
            file_bytes = file.read()
            doc = docx.Document(io.BytesIO(file_bytes))
            normalize_red_runs(doc) # รวมคำสีแดงที่ขาดให้สมบูรณ์
            
            words = extract_all_red_words(doc)
            all_unique_red_words.update(words)
            
            loaded_docs.append({
                "filename": file.filename,
                "doc_obj": doc
            })

        # 2. ให้ AI (Gemini) ทำการจับคู่ข้อมูล (Smart Mapping)
        prompt = f"""
        คุณคือ AI ผู้เชี่ยวชาญด้านการจัดการเอกสาร
        นี่คือรายการ 'คำเดิม' (ข้อความสีแดง) ที่พบในเอกสาร Word:
        {list(all_unique_red_words)}

        และนี่คือ 'ข้อมูลใหม่' ที่ผู้ใช้อัปเดตผ่านตาราง (ข้อมูลเป็น JSON):
        {json.dumps(excel_data, ensure_ascii=False)}

        หน้าที่ของคุณ:
        ให้หาว่า 'คำเดิม' คำไหน ควรถูกแทนที่ด้วยค่าอะไรจาก 'ข้อมูลใหม่' 
        พิจารณาจากบริบทและความหมายให้สอดคล้องกันที่สุด

        ส่งคำตอบกลับมาเป็น JSON Object เท่านั้น โดยให้ key คือ "คำเดิม" และ value คือ "คำใหม่"
        ตัวอย่างเช่น {{"ปี 2568": "2026", "กุมภาพันธ์": "ก.พ."}}
        ห้ามใส่คำอธิบายเพิ่มเติม ห้ามมี Markdown format (เช่น ```json)
        """
        
        response = model.generate_content(prompt)
        result_text = response.text.strip()
        
        # ทำความสะอาดข้อความเผื่อ AI ส่ง markdown ติดมา
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
            
        ai_mapping = json.loads(result_text.strip())
        print("💡 AI Mapping Result:", ai_mapping) # แสดงผลใน Terminal เพื่อให้เราดูว่า AI คิดอะไร

        # 3. นำคู่มือจาก AI ไปแทนที่คำในทุกไฟล์ แล้วแพ็กลง ZIP
        memory_zip = io.BytesIO()
        with zipfile.ZipFile(memory_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            
            for item in loaded_docs:
                doc = item["doc_obj"]
                filename = item["filename"]
                
                # ฟังก์ชันเขียนทับคำ
                def replace_and_recolor(paragraphs):
                    for p in paragraphs:
                        for run in p.runs:
                            if is_reddish(run):
                                original = run.text.strip()
                                # ถ้า AI หาคู่เจอ ให้เปลี่ยนคำและเปลี่ยนเป็นสีดำ
                                if original in ai_mapping and ai_mapping[original] != "":
                                    run.text = run.text.replace(original, str(ai_mapping[original]))
                                    run.font.color.rgb = RGBColor(0, 0, 0) # เปลี่ยนเป็นสีดำ
                                    
                # จัดการย่อหน้าปกติ และ ในตาราง
                replace_and_recolor(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_and_recolor(cell.paragraphs)
                
                # เซฟไฟล์ที่แก้แล้วลงใน RAM แล้วยัดใส่ ZIP
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # เปลี่ยนชื่อไฟล์นิดหน่อยให้รู้ว่าอัปเดตแล้ว
                new_filename = f"อัปเดตแล้ว_{filename}"
                zf.writestr(new_filename, doc_io.read())

        memory_zip.seek(0)
        
        # 4. ส่งไฟล์ ZIP กลับไปให้หน้าเว็บดาวน์โหลด
        return send_file(
            memory_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name='รวมไฟล์เอกสาร_AutoWord.zip'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'เกิดข้อผิดพลาดในการประมวลผล: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)